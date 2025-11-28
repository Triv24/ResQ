import streamlit as st
from langchain_community.document_loaders import PyPDFLoader
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.tools import tool
from langchain_core import documents
from langchain_core.prompts import PromptTemplate
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.output_parsers import StrOutputParser
from docx import Document
import tempfile



# ------------------------------------------------------------------------------------------------------------------------------
# Helper Functions
# ------------------------------------------------------------------------------------------------------------------------------

def init_session():
    if "llm" not in st.session_state:
        st.session_state.llm = ChatGoogleGenerativeAI(
            model = "gemini-2.5-flash",
            google_api_key = st.secrets["GEMINI_API_KEY"]
        )

def load_pdf(file):
    loader = PyPDFLoader(file)
    docs = loader.load()

    full_text = ""
    for doc in docs:
        full_text += doc.page_content

    return full_text

# ------------------------------------------------------------------------------------------------------------------------------
# Tools for Model Question Paper Generator
# ------------------------------------------------------------------------------------------------------------------------------

@tool
def extract_questions(qp):
    """
    This tool will take the content that is loaded from the pdf file of question paper which contains a lot of other things.
    This tool extracts only the questions from the loaded content
    """
    llm = ChatGoogleGenerativeAI(
        model = "gemini-2.5-flash",
        google_api_key = st.secrets["GEMINI_API_KEY"]
    )

    response = llm.invoke(
        """
        please extract all the questions from the contents of a question paper loaded from a pdf file. 
        Output should be clear, concise and to the point.
        Avoid any greetings, just provide the questions.
        start directly from Q1
        ---
        contents : {qp}
        """
    )

    return response.content

@tool
def get_model_questions(question_sets : str) -> str :
    """
    This tool takes the extracted questions in form of string and then uses an llm to analyse the patterns of questions in the question papers and then predicts all the question for model question paper.
    """
    llm = get_llm()

    response = llm.invoke(
        f"""
    You are an expert exam paper analyst and academic question setter with deep experience in identifying patterns across previous examinations.
    ---
    INSTRUCTIONS
    Carefully read and analyze the full contents of the previous question papers provided.
    Identify recurring patterns, frequently tested concepts, question styles, difficulty levels.
    Based on this analysis, estimate and construct a complete model question paper that closely reflects what is most likely to appear in the upcoming exam.
    Ensure the model paper:
    Follows the same structure and format as typical past exams
    Includes appropriate sectioning (Part A, Part B, etc., if relevant)
    ---
    CONTEXT
    Previous question paper content:
    {question_sets}
    ---
    OUTPUT
    A well-structured string, all questions for model question paper that represents the best possible estimation of upcoming exam questions.
    directly start from Q1
    """
    )

    return response.content

@tool
def get_docx(model_questions : str) :
    """
    This tool takes the predicted model questions as an arguement and generates a file path of a word document containing these question.
    """
    document = Document()

    document.add_heading("Model Question Paper", level=2)
    
    document.add_paragraph(model_questions)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        file_path = tmp.name
        document.save(file_path)
    
    return file_path


def get_llm():
    llm = ChatGoogleGenerativeAI(
        model = "gemini-2.5-flash",
        google_api_key = st.secrets["GEMINI_API_KEY"]
    )

    return llm

# --------------------------------------------------------------------------------------------------------------
# Concepts Explaining Assistant Tools 
# --------------------------------------------------------------------------------------------------------------

def get_chunks(text_content):
    """
    This tool takes the loaded content from the pdf of the textbook as arguement and splits them into chunks. Then it converts the text chunks in documents and then returns the list of document chunks.
    """

    # Text Splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size = 1500,
        chunk_overlap = 10
    )

    chunks = text_splitter.split_text(text_content)

    document_chunks = [documents.Document(page_content=chunk) for chunk in chunks]

    return document_chunks

def store_in_db(document_chunks) :
    # Embedding function
    embedder = HuggingFaceEmbeddings(
        model_name = "sentence-transformers/all-MiniLM-L6-v2",
    model_kwargs = {
        "device": "cpu"
    },
    encode_kwargs = {
        'normalize_embeddings': True
    }
    )

    # Vector store
    db = Chroma(
        embedding_function=embedder,
        persist_directory="my_knowledge_base",
        collection_name="textbook_data"
    )

    db.add_documents(document_chunks)

    return db

@tool
def get_relevant_chunks(query : str):
    """
    This tool takes the query as arguements and then returns the top 5 relevant results of similarity search from the vector store. The top 5 relevant chunks are returned in the form of plain string.
    """
    embedder = HuggingFaceEmbeddings(
        model_name = "sentence-transformers/all-MiniLM-L6-v2",
    model_kwargs = {
        "device": "cpu"
    },
    encode_kwargs = {
        'normalize_embeddings': True
    }
    )

    db = Chroma(
        embedding_function=embedder,
        persist_directory="my_knowledge_base",
        collection_name="textbook_data"
    )

    relevant_chunks = ""

    for chunk in db.similarity_search(query, k=5):
        relevant_chunks += chunk.page_content

    return relevant_chunks

@tool
def explain_concept(relevant_chunks : str, query : str):
    """
    This tool takes relevant chunks from the vector store and the query as arguements and return the generated explaination of the concept mentioned in query with suitable examples.
    """

    llm = ChatGoogleGenerativeAI(
        model = "gemini-2.5-flash",
        google_api_key = st.secrets["GEMINI_API_KEY"]
    )

    parser = StrOutputParser()

    prompt = PromptTemplate(
        template="""
        You are an expert educator, curriculum designer, and AI explanation specialist trained to teach complex topics in simple, clear, student-friendly language, using real-life examples and context-aware reasoning.

INSTRUCTIONS (READ CAREFULLY BEFORE RESPONDING)

Use information found in the relevant chunks provided:
you can add some extra info only if needed
{relevant_chunks}

You must explain the concept asked in the user query:
{query}

Your explanation must follow these rules:

Use very simple language, suitable for a high-school or first-year college student.

Break the explanation into clear sections and step-by-step reasoning.

Add real-life analogies, examples, or stories that make the concept easy to understand.

Avoid jargon. If a technical term is required, define it clearly.

Stay fully grounded in the content of the provided chunks; do not invent unsupported information.

Always end with a short “Why this concept matters” section.

If any part of the query cannot be answered from the chunks, you must clearly state which part is missing.

output format should be string of answer in markdown.
        """,
        input_variables=["query", "relevant_chunks"]
    )

    chain = prompt | llm | parser
    
    response = chain.invoke(
        {
            "query" : query,
            "relevant_chunks" : relevant_chunks
        }
    )

    return response

    
# --------------------------------------------------------------------------------------------------------------
# Assignments Helper
# --------------------------------------------------------------------------------------------------------------

